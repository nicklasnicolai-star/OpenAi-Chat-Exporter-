import os
import json
from docx import Document
from docx.shared import Pt
import re

# === CONFIG ===
INPUT_JSON = "conversations.json"
OUTPUT_DIR = "chats_output"

# === SETUP ===
os.makedirs(OUTPUT_DIR, exist_ok=True)

with open(INPUT_JSON, "r", encoding="utf-8") as f:
    conversations = json.load(f)

def extract_messages(conversation):
    messages = []
    mapping = conversation.get("mapping", {})
    root_id = next((k for k, v in mapping.items() if v.get("parent") is None), None)
    queue = [root_id] if root_id else []

    while queue:
        current = queue.pop(0)
        node = mapping.get(current, {})
        msg = node.get("message")
        if msg and msg.get("author") and msg.get("content") and msg["content"].get("parts"):
            author = msg["author"]["role"]
            parts = msg["content"]["parts"]
            full_text = "\n".join(
                part if isinstance(part, str) else part.get("text", "")
                for part in parts if part
            ).strip()
            if full_text:
                messages.append((author, full_text))
        # Füge Kinder in richtiger Reihenfolge hinzu
        children = node.get("children", [])
        queue.extend(children)
    return messages

def clean_text(s):
    # Entferne alle nicht-XML-kompatiblen Steuerzeichen (außer \n, \t)
    return re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F]', '', s)

def save_docx(title, messages, index):
    safe_title = f"{index:04d}_{title.replace('/', '_')[:60]}.docx"
    doc_path = os.path.join(OUTPUT_DIR, safe_title)
    doc = Document()

    doc.add_heading(title.strip(), level=1)

    for author, raw_text in messages:
        role = "User" if author == "user" else "ChatGPT"
        para = doc.add_paragraph()
        para.add_run(f"{role}:\n").bold = True

        # ✅ Clean text
        clean = clean_text(raw_text)
        para.add_run(clean).font.size = Pt(12)
        doc.add_paragraph()  # Leerzeile

    doc.save(doc_path)

# === MAIN ===
for i, conv in enumerate(conversations):
    title = conv.get("title", f"Untitled_{i}")
    messages = extract_messages(conv)
    if messages:
        save_docx(title, messages, i)

print(f"Fertig! Exportiert nach: {OUTPUT_DIR}")